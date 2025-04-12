import streamlit as st
import os
from database import supabase_client as supabase
from PyPDF2 import PdfReader
from google.generativeai import configure, GenerativeModel
from dotenv import load_dotenv
from docx import Document
import re

load_dotenv()

# Configure Gemini-1.5-Pro API
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
configure(api_key=GEMINI_API_KEY)
model = GenerativeModel("gemini-1.5-pro")


def fetch_document_content(file_name):
    """Fetches and reads the selected document from session state or Supabase Storage."""
    if "selected_document_text" in st.session_state:
        return st.session_state["selected_document_text"]

    user_display_name = st.session_state["username"]
    selected_chat = st.session_state.get("selected_chat")

    if not selected_chat:
        st.sidebar.error("Please select a chat first.")
        return None

    bucket_name = "user-documents"
    file_path = f"{user_display_name}/{selected_chat}/{file_name}"

    try:
        response = supabase.storage.from_(bucket_name).download(file_path)
        if response:
            document_text = response.decode("utf-8")
            st.session_state["selected_document_text"] = document_text
            return document_text
        else:
            st.error("Failed to retrieve the document.")
            return None
    except Exception as e:
        st.error(f"Error retrieving document: {e}")
        return None


def analyze_notes(content, user_prompt):
    """Uses Gemini AI to analyze and enhance notes."""
    prompt = (
        f"Analyze and enhance the following notes for better learning. "
        f"User request: {user_prompt}\n\n{content}"
    )
    try:
        response = model.generate_content(prompt)
        return response.text if response else "AI analysis failed."
    except Exception as e:
        st.error(f"Error in AI analysis: {e}")
        return None


def create_docx(text):
    """Generates a properly formatted DOCX file from the enhanced notes."""
    doc = Document()

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            doc.add_paragraph("")  # Preserve blank lines
            continue

        # Headings (Markdown-style)
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue

        # Bullet Points
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="ListBullet")
            continue

        # Numbered List
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(line, style="ListNumber")
            continue

        # Handle bold and italic formatting inside paragraphs
        para = doc.add_paragraph()
        pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")  # Match **bold** and *italic*
        matches = pattern.split(line)  # Split line into formatted segments

        for part in matches:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])  # Remove ** and apply bold
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])  # Remove * and apply italic
                run.italic = True
            else:
                para.add_run(part)  # Add normal text

    doc_path = "Enhanced_Notes.docx"
    doc.save(doc_path)

    return doc_path

def notes_page():
    """Displays the AI-Enhanced Notes page in Streamlit."""
    st.title("📑 AI-Enhanced Notes")
    st.write("AI will enhance your selected document for better learning.")

    file_content = st.session_state.get("selected_document_text")

    if not file_content:
        st.warning("⚠️ No document content available. Please upload or select a document.")
        return

    text = ""
    try:
        if isinstance(file_content, bytes):
            from io import BytesIO
            pdf_reader = PdfReader(BytesIO(file_content))
            text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        else:
            text = file_content
    except Exception as e:
        st.error(f"Error processing document: {e}")
        return

    st.text_area("📄 Document Content", text, height=300, disabled=True)
    user_prompt = st.text_area("✍️ Specify your learning focus",
                               placeholder="Summarize key concepts, explain acronyms, etc.")

    if st.button("🧠 Enhance Notes"):
        enhanced_notes = analyze_notes(text, user_prompt)
        st.session_state["enhanced_notes"] = enhanced_notes
        st.markdown(enhanced_notes, unsafe_allow_html=True)

        docx_path = create_docx(enhanced_notes)
        with open(docx_path, "rb") as docx_file:
            st.download_button(label="📥 Download Enhanced Notes (DOCX)", data=docx_file,
                               file_name="Enhanced_Notes.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
